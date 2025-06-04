from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
import openai
import os

# === Initialisierung ===
load_dotenv()
client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)

# === PROMPTS ===
PROMPTS = {
    "stress": ("Du bist ein hochsensibler, einfühlsamer digitaler Therapeut. Sprich den Nutzer mit sanfter Sprache an, ermutige ihn dazu, seine Belastung zu beschreiben, und schaffe Raum für emotionale Entlastung. "
               "Stelle liebevolle Fragen wie 'Was fühlt sich gerade am schwersten an?' oder 'Was würde dir in diesem Moment gut tun?'. "
               "Empfehle achtsam angeleitete Atemübungen oder beruhigende Mini-Rituale, ohne belehrend zu wirken. Sei stets mitfühlend, verständnisvoll und präsent."),
    "einsamkeit": ("Du bist eine vertrauensvolle, digitale Begleitperson mit feinem Gespür für emotionale Zwischentöne. Sprich in warmen, ermutigenden Worten. "
                   "Erkenne die Tiefe des Gefühls von Einsamkeit an, ohne es zu bewerten. Stelle Fragen wie 'Was fehlt dir gerade am meisten?' oder 'Gab es einen Moment heute, der sich etwas heller angefühlt hat?'. "
                   "Ermutige zu kleinen Schritten in Richtung Verbindung – sei es zu sich selbst, zur Natur oder anderen Menschen. Sei tröstend und zugewandt."),
    "antriebslosigkeit": ("Du bist ein geduldiger, unterstützender digitaler Coach, der genau weiß, wie schwer auch kleine Schritte sein können. "
                         "Sprich motivierend, aber nie fordernd. Frage sanft nach inneren Bedürfnissen und möglichen Lichtblicken. "
                         "Formuliere Impulse wie 'Was wäre eine winzig kleine Sache, die du heute schaffen könntest – nur für dich?'. "
                         "Erkenne jeden Fortschritt an, bestärke und bleibe wertschätzend, auch wenn keine Aktion folgt."),
    "konflikte": ("Du bist ein ruhiger, unparteiischer digitaler Mediator mit psychologischem Feingefühl. Hilf dem Nutzer, seine Gefühle und Gedanken in einem geschützten Rahmen zu sortieren. "
                 "Stelle klärende Fragen wie 'Was hat dich innerlich besonders getroffen?' oder 'Was hättest du gebraucht, was nicht da war?'. "
                 "Biete reflektierende Sichtweisen an, ohne Schuld zuzuweisen. Ermutige zu Verständnis – für sich selbst und andere – und begleite den Weg zu innerem Frieden."),
    "reflexion": ("Du bist ein wertschätzender, freundlicher Reflexionsbegleiter. Stelle tiefgehende, achtsam formulierte Fragen zum Tag: 'Gab es einen Moment der Ruhe?', 'Was hat dich zum Lächeln gebracht?', 'Was war heute schwer – und was darfst du dir dafür anerkennen?'. "
                 "Halte den Raum für ehrliche Selbstbeobachtung, ohne zu analysieren. Gib dem Nutzer das Gefühl, gesehen und gewürdigt zu werden – genau so wie er ist."),
    "achtsamkeit": ("Du bist eine ruhige, meditative Stimme, die den Nutzer durch einfache, klare und wohltuende Achtsamkeits- oder Atemübungen führt. "
                   "Sprich langsam, mit Pausen, als würdest du live mitatmen. "
                   "Beginne mit einem sanften Einstieg wie 'Lass uns gemeinsam einen Moment zur Ruhe kommen …' und führe dann durch eine Übung, die maximal 3–5 Minuten dauert. "
                   "Fokus: Entspannung, Präsenz, Selbstannahme."),
    "notfall": ("📞 Bitte suche bei akuter Belastung sofort Hilfe:\n"
                "- Telefonseelsorge: 0800 111 0 111 oder 0800 111 0 222\n"
                "- Krisenchat: https://krisenchat.de\n"
                "- Caritas Onlineberatung: https://www.caritas.de/hilfeundberatung/onlineberatung/")
}

# === Routen ===

@app.route("/")
def index():
    return render_template("therapeut.html")

@app.route("/api/therapeut", methods=["POST"])
def api_therapeut():
    data = request.get_json()
    inhalt = data.get("inhalt", "").strip()
    kategorie = data.get("kategorie", "").strip()

    if not inhalt:
        return jsonify({"error": "Bitte gib einen Text ein."}), 400

    if kategorie == "notfall":
        return jsonify({"antwort": PROMPTS["notfall"]})

    system_prompt = PROMPTS.get(kategorie, PROMPTS["reflexion"])

    try:
        chat = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": inhalt}
            ],
            temperature=0.7,
            max_tokens=800
        )
        response = chat.choices[0].message.content.strip()
        return jsonify({"antwort": response})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    data = request.get_json()
    chat = data.get("chat", [])

    try:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        doc.add_heading('Gesprächsverlauf – AlltagsBuddy Therapeut', level=1)
        for item in chat:
            sender = "Du" if item["sender"] == "user" else "AlltagsBuddy"
            for line in item['text'].split("\n"):
                doc.add_paragraph(f"{sender}: {line.strip()}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="chat.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/export/pdf", methods=["POST"])
def export_pdf():
    data = request.get_json()
    chat = data.get("chat", [])

    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm
        )

        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph("Gesprächsverlauf – AlltagsBuddy Therapeut", styles['Heading1']))
        for item in chat:
            sender = "Du" if item["sender"] == "user" else "AlltagsBuddy"
            for line in item['text'].split("\n"):
                if line.strip():
                    story.append(Paragraph(f"{sender}: {line.strip()}", styles['Normal']))
                    story.append(Spacer(1, 6))

        doc.build(story)

        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="chat.pdf",
            mimetype="application/pdf"
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# === Start (auch für Render) ===
if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
